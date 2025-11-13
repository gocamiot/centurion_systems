import csv
import json
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.contrib.auth.decorators import login_required
from apps.tables.utils import (
    software_filter, 
    same_key_filter, 
    common_unique_filter,
    common_count_filter,
    get_user_id,
    user_filter_common_func,
    create_page_items,
    create_hide_show_filter,
    common_snapshot_filter, get_model_fields,
    common_date_filter, common_integer_filter, common_float_filter
)
from apps.tables.models import (
    UserFilter, 
    PageItems, 
    HideShowFilter, 
    DateRangeFilter, 
    IntRangeFilter, 
    FloatRangeFilter, 
    ModelChoices, 
    Finding,
    ActionStatus,
    EmailActionStatus,
    TableDropdownSubItem,
    FindingAction,
    DependentDropdown
)
from django.http import HttpResponse
from django.views import View
from django.http import JsonResponse
from datetime import datetime
from django.utils import timezone
from apps.tables.forms import FindingForm, ActionForm
from django.urls import reverse
from django.contrib.auth import get_user_model
from django.core.mail import send_mail
from django.conf import settings
from django.contrib import messages
from django.http import HttpResponse
from django.views import View
from django.http import JsonResponse
from datetime import datetime
from django.db.models import F
from apps.common.models import SavedFilter, FieldType

User = get_user_model()

map_field_names = getattr(settings, 'MAP_FIELD_NAMES')
reverse_map_field_names = {v: k for k, v in map_field_names.items()}

# Create your views here.

def create_finding_filter(request):
    if request.method == 'POST':
        user_filter_common_func(request, 'Finding', 'FINDING_VIEW')
    return redirect(request.META.get('HTTP_REFERER'))


def create_finding_page_items(request):
    if request.method == 'POST':
        create_page_items(request, 'FINDING_VIEW')
        return redirect(request.META.get('HTTP_REFERER'))

def create_finding_hide_show_filter(request):
    if request.method == "POST":
        response_data = create_hide_show_filter(request, 'FINDING_VIEW')
        return JsonResponse(response_data)

    return JsonResponse({'error': 'Invalid request'}, status=400)
    
def delete_finding_view_filter(request, id):
    filter_instance = UserFilter.objects.get(id=id, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
    filter_instance.delete()
    return redirect(reverse('finding'))

def delete_finding_daterange_filter(request, id):
    date_filter = DateRangeFilter.objects.get(id=id, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
    date_filter.delete()
    return redirect(reverse('finding'))

def delete_finding_intrange_filter(request, id):
    int_filter = IntRangeFilter.objects.get(id=id, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
    int_filter.delete()
    return redirect(reverse('finding'))

def delete_finding_floatrange_filter(request, id):
    float_filter = FloatRangeFilter.objects.get(id=id, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
    float_filter.delete()
    return redirect(reverse('finding'))

@login_required(login_url='/users/signin/')
def finding(request, companies=None):
    db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]

    # Snapshot
    latest_snapshot, summary, snapshot_filter, snapshots = common_snapshot_filter(request, 'finding')

    field_names = []
    for field_name in db_field_names:
        mapped_field_name = map_field_names.get(field_name, field_name)
        fields, created = HideShowFilter.objects.get_or_create(key=mapped_field_name, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
        field_names.append(fields)

    page_items = PageItems.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW).last()
    items = 25
    if page_items:
        items = page_items.items_per_page

    filter_string = {}
    pre_filter_string = {}
    pre_filter_string['email_action_status__exact'] = 'OPEN'
    if companies:
        pre_filter_string['companies__exact'] = companies

    if finding_id := request.GET.get('finding_id'):
        SavedFilter.objects.filter( userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_name='id').delete()
        SavedFilter.objects.get_or_create(
            userID=get_user_id(request), 
            parent=ModelChoices.FINDING_VIEW,
            field_name='id',
            field_type=FieldType.TEXT,
            is_hidden=True,
            defaults={
                'value_start': finding_id
            }
        )
    else:
        SavedFilter.objects.filter( userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_name='id').delete()

    filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.TEXT)
    combined_q_objects, user_unique_filter, user_query_conditions, user_count_filters = same_key_filter(filter_instance, return_count_filters=True)

    # for date range
    date_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.DATE)
    date_string, date_unique_filter, date_query_conditions, date_count_filters = common_date_filter(date_filter_instance, return_count_filters=True)
    filter_string.update(date_string)
    
    # for integer range
    int_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.INT)
    int_string, int_unique_filter, int_query_conditions, int_count_filters = common_integer_filter(int_filter_instance, return_count_filters=True)
    filter_string.update(int_string)

    # for float range
    float_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.FLOAT)
    float_string, float_unique_filter, float_query_conditions, float_count_filters = common_float_filter(float_filter_instance, return_count_filters=True)
    filter_string.update(float_string)

    order_by = request.GET.get('order_by', 'id')
    base_queryset = Finding.objects.filter(parent=None).filter(action_status=ActionStatus.IS_ACTIVE).filter(**filter_string).filter(**pre_filter_string).filter(combined_q_objects).filter(**snapshot_filter)   
    queryset = base_queryset

    if user_query_conditions:
        queryset = queryset.filter(user_query_conditions)
    if date_query_conditions:
        queryset = queryset.filter(date_query_conditions)
    if int_query_conditions:
        queryset = queryset.filter(int_query_conditions)
    if float_query_conditions:
        queryset = queryset.filter(float_query_conditions)

    if user_count_filters:
        queryset = common_count_filter(user_count_filters, base_queryset, queryset, db_field_names)
    elif date_count_filters:
        queryset = common_count_filter(date_count_filters, base_queryset, queryset, db_field_names)
    elif int_count_filters:
        queryset = common_count_filter(int_count_filters, base_queryset, queryset, db_field_names)
    elif float_count_filters:
        queryset = common_count_filter(float_count_filters, base_queryset, queryset, db_field_names)
    else:
        if order_by == 'count' or order_by == '-count':
            order_by = 'ID'

    if user_unique_filter:
        queryset = common_unique_filter(request, user_unique_filter, queryset, snapshot_filter, 'tables_finding')
    if date_unique_filter:
        queryset = common_unique_filter(request, date_unique_filter, queryset, snapshot_filter, 'tables_finding')
    if int_unique_filter:
        queryset = common_unique_filter(request, int_unique_filter, queryset, snapshot_filter, 'tables_finding')
    if float_unique_filter:
        queryset = common_unique_filter(request, float_unique_filter, queryset, snapshot_filter, 'tables_finding')

    # Order by
    if order_by == 'count' and 'count' not in db_field_names:
        queryset = queryset.order_by(F('count').desc(nulls_last=True))
    else:
        queryset = queryset.order_by(order_by)

    software_list = software_filter(request, queryset, db_field_names)

    page = request.GET.get('page', 1)
    paginator = Paginator(software_list, items)
    
    try:
        softwares = paginator.page(page)
    except PageNotAnInteger:
        return redirect('finding')
    except EmptyPage:
        return redirect('finding') 

    read_only_fields = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by', 'created_at', 'updated_at', 'version_control', 'created_by', 'updated_by',)
    pre_column = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions','itgc_focus_area','risk','control',)
    compulsory_fields = ('description','recommendation','status',)
    richtext_fields = ('description', 'recommendation', )
    dropdown_fields = ('companies', 'itgc_categories', 'itgc_questions', )
    dropdown_fields = [
        map_field_names[field] if field in map_field_names else field
        for field in dropdown_fields
    ]

    # date_picker_fields = ('DataAge','AD_Account_Created','Password_Last_Set')
    finding_tabs = DependentDropdown.objects.filter(parent__title='Finding').values_list('title', flat=True)
    companies_dropdown = DependentDropdown.objects.filter(parent__title='Finding')
    finding_categories = TableDropdownSubItem.objects.filter(item__item='Finding_Cat').values_list('subitem', flat=True)
    finding_questions = TableDropdownSubItem.objects.filter(item__item='Finding_Question').values_list('subitem', flat=True)
    action_type = TableDropdownSubItem.objects.filter(item__item='Action_Type').values_list('subitem', flat=True)

    COMBOS = {}
    COMBOS['action_to'] = User.objects.exclude(email="").values_list('email', flat=True)
    COMBOS['email_action_status'] = [choice[0] for choice in EmailActionStatus.choices]
    COMBOS['companies'] = finding_tabs
    COMBOS['itgc_categories'] = finding_categories
    COMBOS['itgc_questions'] = finding_questions
    COMBOS['action_type'] = action_type   
  
    form = FindingForm()
    action_form = ActionForm(action_status=False)

    fields = get_model_fields('Finding', pre_column)
    saved_filters = list(SavedFilter.objects.filter(
        userID=request.user.id, parent=ModelChoices.FINDING_VIEW
    ).values())
    for filter in saved_filters:
        if 'created_at' in filter:
            filter['created_at'] = filter['created_at'].isoformat()
    saved_filters_json = json.dumps(saved_filters)

    updated_db_field_names = [
        map_field_names[field] if field in map_field_names else field
        for field in db_field_names
    ]

    context = {
        'segment'  : 'finding',
        'parent'   : 'audit_report',
        'title'    : 'All Actions',
        'softwares' : softwares,
        'field_names': field_names,
        'db_field_names': updated_db_field_names,
        'filter_instance': filter_instance,
        'date_filter_instance': date_filter_instance,
        'items': items,
        'read_only_fields': read_only_fields,
        'total_items': Finding.objects.filter(parent=None).filter(action_status=ActionStatus.IS_ACTIVE).filter(**pre_filter_string).filter(**snapshot_filter).count(),
        'pre_column': pre_column,
        'COMBOS': COMBOS,
        'date_picker_fields': Finding.date_fields_to_convert,
        'compulsory_fields': compulsory_fields,
        'richtext_fields': richtext_fields,
        'dropdown_fields': dropdown_fields,
        'fav_name': 'Finding',
        'int_filter_instance': int_filter_instance,
        'int_fields': Finding.integer_fields,
        'float_filter_instance': float_filter_instance,
        'float_fields': Finding.float_fields,
        'form': form,
        'action_form': action_form,
        'companies': companies,
        'companies_dropdown': companies_dropdown,
        'fields': fields,
        'saved_filters': saved_filters_json,
        'unique_filter': list(set(user_unique_filter + date_unique_filter + int_unique_filter + float_unique_filter)),
        'table_name': 'finding',
        'finding_url': request.build_absolute_uri(reverse('finding')),

        # snapshots
        'snapshots': snapshots,
        'latest_snapshot': latest_snapshot,
        'selected_snapshot': summary
    }
    
    return render(request, 'apps/finding.html', context)


@login_required(login_url='/users/signin/')
def my_finding(request):
    action_status = request.GET.get('action_status')
    action_selected = action_status if action_status else 'OPEN'
    db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]
    updated_db_field_names = [
        map_field_names[field] if field in map_field_names else field
        for field in db_field_names
    ]
    action_filter = {}
    action_filter['status__exact'] = 'OPEN'
    if action_status:
        if not action_status == 'all':
            action_filter['status__exact'] = action_status
        else:
            action_filter = {}

    finding_ids = FindingAction.objects.filter(**action_filter).filter(action_to=request.user.email).values_list('finding__id', flat=True)

    # Snapshot
    latest_snapshot, summary, snapshot_filter, snapshots = common_snapshot_filter(request, 'finding')

    field_names = []
    for field_name in db_field_names:
        mapped_field_name = map_field_names.get(field_name, field_name)
        fields, created = HideShowFilter.objects.get_or_create(key=mapped_field_name, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
        field_names.append(fields)

    page_items = PageItems.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW).last()
    items = 25
    if page_items:
        items = page_items.items_per_page

    filter_string = {}
    pre_filter_string = {}
    pre_filter_string['email_action_status__exact'] = 'OPEN'

    filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.TEXT)
    combined_q_objects, user_unique_filter, user_query_conditions, user_count_filters = same_key_filter(filter_instance, return_count_filters=True)

    # for date range
    date_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.DATE)
    date_string, date_unique_filter, date_query_conditions, date_count_filters = common_date_filter(date_filter_instance, return_count_filters=True)
    filter_string.update(date_string)
    
    # for integer range
    int_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.INT)
    int_string, int_unique_filter, int_query_conditions, int_count_filters = common_integer_filter(int_filter_instance, return_count_filters=True)
    filter_string.update(int_string)

    # for float range
    float_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.FLOAT)
    float_string, float_unique_filter, float_query_conditions, float_count_filters = common_float_filter(float_filter_instance, return_count_filters=True)
    filter_string.update(float_string)

    order_by = request.GET.get('order_by', 'id')
    base_queryset = Finding.objects.filter(id__in=finding_ids).filter(parent=None).filter(action_status=ActionStatus.IS_ACTIVE).filter(**filter_string).filter(**pre_filter_string).filter(combined_q_objects).filter(**snapshot_filter)
    queryset = base_queryset

    if user_query_conditions:
        queryset = queryset.filter(user_query_conditions)
    if date_query_conditions:
        queryset = queryset.filter(date_query_conditions)
    if int_query_conditions:
        queryset = queryset.filter(int_query_conditions)
    if float_query_conditions:
        queryset = queryset.filter(float_query_conditions)
    
    if user_count_filters:
        queryset = common_count_filter(user_count_filters, base_queryset, queryset, db_field_names)
    elif date_count_filters:
        queryset = common_count_filter(date_count_filters, base_queryset, queryset, db_field_names)
    elif int_count_filters:
        queryset = common_count_filter(int_count_filters, base_queryset, queryset, db_field_names)
    elif float_count_filters:
        queryset = common_count_filter(float_count_filters, base_queryset, queryset, db_field_names)
    else:
        if order_by == 'count' or order_by == '-count':
            order_by = 'ID'

    if user_unique_filter:
        queryset = common_unique_filter(request, user_unique_filter, queryset, snapshot_filter, 'tables_finding')
    if date_unique_filter:
        queryset = common_unique_filter(request, date_unique_filter, queryset, snapshot_filter, 'tables_finding')
    if int_unique_filter:
        queryset = common_unique_filter(request, int_unique_filter, queryset, snapshot_filter, 'tables_finding')
    if float_unique_filter:
        queryset = common_unique_filter(request, float_unique_filter, queryset, snapshot_filter, 'tables_finding')

    # Order by
    if order_by == 'count' and 'count' not in db_field_names:
        queryset = queryset.order_by(F('count').desc(nulls_last=True))
    else:
        queryset = queryset.order_by(order_by)

    software_list = software_filter(request, queryset, db_field_names)

    page = request.GET.get('page', 1)
    paginator = Paginator(software_list, items)
    
    try:
        softwares = paginator.page(page)
    except PageNotAnInteger:
        return redirect('finding')
    except EmptyPage:
        return redirect('finding') 

    read_only_fields = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by', 'created_at', 'updated_at', 'version_control', 'created_by', 'updated_by',)
    pre_column = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions','itgc_focus_area','risk','control',)
    compulsory_fields = ('status',)
    richtext_fields = ('description', 'recommendation', )
    dropdown_fields = ('companies', 'itgc_categories', 'itgc_questions', )
    dropdown_fields = [
        map_field_names[field] if field in map_field_names else field
        for field in dropdown_fields
    ]

    finding_tabs = DependentDropdown.objects.filter(parent__title='Finding').values_list('title', flat=True)
    finding_categories = TableDropdownSubItem.objects.filter(item__item='Finding_Cat').values_list('subitem', flat=True)
    finding_questions = TableDropdownSubItem.objects.filter(item__item='Finding_Question').values_list('subitem', flat=True)
    action_type = TableDropdownSubItem.objects.filter(item__item='Action_Type').values_list('subitem', flat=True)
    
    COMBOS = {}
    COMBOS['action_to'] = User.objects.exclude(email="").values_list('email', flat=True).distinct()
    COMBOS['email_action_status'] = [choice[0] for choice in EmailActionStatus.choices]
    form = FindingForm()
    action_form = ActionForm(action_status=False)
    COMBOS['companies'] = finding_tabs
    COMBOS['itgc_categories'] = finding_categories
    COMBOS['itgc_questions'] = finding_questions
    COMBOS['action_type'] = action_type  

    fields = get_model_fields('Finding', pre_column)
    saved_filters = list(SavedFilter.objects.filter(
        userID=request.user.id, parent=ModelChoices.FINDING_VIEW
    ).values())
    for filter in saved_filters:
        if 'created_at' in filter:
            filter['created_at'] = filter['created_at'].isoformat()
    saved_filters_json = json.dumps(saved_filters)

    context = {
        'segment'  : 'finding',
        'parent'   : 'audit_report',
        'title'    : 'My Actions',
        'softwares' : softwares,
        'field_names': field_names,
        'db_field_names': updated_db_field_names,
        'filter_instance': filter_instance,
        'date_filter_instance': date_filter_instance,
        'items': items,
        'read_only_fields': read_only_fields,
        'total_items': Finding.objects.filter(id__in=finding_ids).filter(parent=None).filter(action_status=ActionStatus.IS_ACTIVE).filter(**pre_filter_string).filter(**snapshot_filter).count(),
        'pre_column': pre_column,
        'COMBOS': COMBOS,
        'date_picker_fields': Finding.date_fields_to_convert,
        'compulsory_fields': compulsory_fields,
        'richtext_fields': richtext_fields,
        'dropdown_fields': dropdown_fields,
        'fav_name': 'Finding',
        'int_filter_instance': int_filter_instance,
        'int_fields': Finding.integer_fields,
        'float_filter_instance': float_filter_instance,
        'float_fields': Finding.float_fields,
        'form': form,
        'action_form': action_form,
        'action_selected': action_selected,
        'fields': fields,
        'saved_filters': saved_filters_json,
        'unique_filter': list(set(user_unique_filter + date_unique_filter + int_unique_filter + float_unique_filter)),
        'table_name': 'finding',
        'finding_url': request.build_absolute_uri(reverse('finding')),

        # snapshots
        'snapshots': snapshots,
        'latest_snapshot': latest_snapshot,
        'selected_snapshot': summary
    }
    
    return render(request, 'apps/my-finding.html', context)


def send_action_email(finding, finding_action):
    subject = "BDO GRC Server - Pending Action"

    message = f"<p>The following action has been assigned to you in the BDO GRC Server.</p>"

    def get_dropdown_title(value):
        try:
            value = int(value)
            dropdown = DependentDropdown.objects.filter(pk=value).first()
            if dropdown:
                return dropdown.title
        except:
            return ''
        return ''

    if finding.companies:
        company = get_dropdown_title(finding.companies)
        message += f"<br><strong>{map_field_names['companies']}: </strong> {company}" 

    if finding.itgc_categories:
        itgc_category = get_dropdown_title(finding.itgc_categories)
        message += f"<br><strong>{map_field_names['itgc_categories']}: </strong> {itgc_category}"    

    if finding.itgc_questions:
        itgc_question = get_dropdown_title(finding.itgc_questions)
        message += f"<br><strong>{map_field_names['itgc_questions']}: </strong> {itgc_question}" 

    if finding.action_note:
        message += f"<br><strong>Action Note: </strong> {finding.action_note}"
    if finding.action_deadline:
        message += f"<br><strong>Deadline Date: </strong> {finding.action_deadline}"
    
    message += f"<br><br><strong>Description:</strong> " + finding.description.html

    if finding.recommendation:
        message += f"<strong>Recommendation: </strong>" + finding.recommendation.html   

    current_page_link = f"{settings.SITE_URL}/tables/my-actions/"

    message += f"<br><br><a href='{current_page_link}'>View more details about this action</a>"            

    message += f"<br><br><p>© 2024 | <a href=\"https://www.bdo.global/en-gb/home\">BDO</a> BDO GRC Server - All rights reserved<p>"   
               
    try:
        send_mail(
            subject=subject,
            message="",
            html_message=message,
            from_email=getattr(settings, 'DEFAULT_FROM_EMAIL'),
            recipient_list=[finding_action.action_to],
        )
    except: pass

@login_required(login_url='/users/signin/')
def create_finding(request):
    if request.method == 'POST':
        software_data = {}
        for attribute, value in request.POST.items():
            if attribute == 'csrfmiddlewaretoken' or attribute == 'action_to':
                continue
            
            if value == '':
                if attribute in Finding.integer_fields:
                    value = None
                elif attribute in Finding.float_fields:
                    value = None
                else:
                    continue

            if attribute not in Finding.date_fields_to_convert:
                if attribute in Finding.integer_fields:
                    software_data[attribute] = int(value) if value else None
                elif attribute in Finding.float_fields:
                    software_data[attribute] = float(value) if value else None
                else:
                    software_data[attribute] = value if value else ''
            else:
                unix_time = datetime.strptime(datetime.now().strftime("%Y-%m-%dT%H:%M"), "%Y-%m-%dT%H:%M").timestamp()
                if value:
                    unix_time = datetime.strptime(value, "%Y-%m-%dT%H:%M").timestamp()
                software_data[attribute] = unix_time

        software_data['user'] = request.user
        finding = Finding.objects.create(
            created_by=request.user.email if request.user.email else request.user.username,
            created_at=timezone.now(),
            #updated_at=timezone.now(),
            **software_data
        )

        action_to = request.POST.getlist('action_to')
        for email in action_to:
            finding_action, created = FindingAction.objects.get_or_create(
                finding=finding,
                action_to=email
            )
            if finding_action.action_to:
                send_action_email(finding, finding_action)


    return redirect(request.META.get('HTTP_REFERER'))

@login_required(login_url='/users/signin/')
def post_request_handling(request, form):
    form.save()
    return redirect(request.META.get('HTTP_REFERER'))

@login_required(login_url='/users/signin/')
def delete_finding(request, id):
    software = Finding.objects.get(id=id)
    software.action_status = ActionStatus.DELETED
    software.deleted_at = timezone.now()
    software.deleted_by = request.user.email if request.user.email else request.user.username
    software.save()
    return redirect(request.META.get('HTTP_REFERER'))


@login_required(login_url='/users/signin/')
def get_sub_items(request, pk):
    try:
        parent_item = DependentDropdown.objects.get(pk=pk)
        sub_items = DependentDropdown.objects.filter(parent=parent_item, featured=False)
        sub_items_data = [{"id": item.pk, "title": item.title} for item in sub_items]
        return JsonResponse({"sub_items": sub_items_data})
    except DependentDropdown.DoesNotExist:
        return JsonResponse({"error": "Parent item not found"}, status=404)

@login_required(login_url='/users/signin/')
def update_finding(request, id):
    software = Finding.objects.get(id=id)
    if request.method == 'POST':
        description = request.POST.get('description', '').strip()
        description_html = None
        if description:
            description_json = json.loads(description)
            description_html = description_json.get('html', '').strip()

        if not description_html or description_html == '<p><br></p>':
            messages.error(request, 'Description cannot be empty.')
            return redirect(request.META.get('HTTP_REFERER'))

        # fields = software._meta.get_fields()
        fields = [field.name for field in software._meta.get_fields() if not field.is_relation]

        fields_and_values = {}
        for field in fields:
            try:
                field_value = getattr(software, field)
                fields_and_values[field] = field_value
            except:
                continue

        fields_and_values.pop('id')
        # fields_and_values.pop('parent')
        fields_and_values.pop('created_at')
        fields_and_values.pop('created_by')
        fields_and_values.pop('updated_at')
        fields_and_values.pop('updated_by')
        fields_and_values.pop('version_control')

        if request.method == 'POST':
            for attribute, value in request.POST.items():
                if attribute == 'csrfmiddlewaretoken':
                    continue

                if value == '':
                    if attribute in Finding.integer_fields:
                        value = None
                    elif attribute in Finding.float_fields:
                        value = None
                    else:
                        continue

                if not attribute in Finding.date_fields_to_convert:
                    setattr(software, attribute, value)
                else:
                    unix_time = setattr(software, attribute, value)
                    possible_formats = ['%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M', '%Y-%m-%d %H:%M:%S', '%Y-%m-%d']
                    for format_str in possible_formats:
                        try:
                            if value:
                                unix_time = datetime.strptime(value, format_str).timestamp()
                                setattr(software, attribute, unix_time)
                                break
                        except ValueError:
                            pass 
            
            if software.action_status == 'DELETED':
                software.deleted_by = request.user.email if request.user.email else request.user.username
                software.deleted_at = timezone.now()
            else:
                software.updated_by = request.user.email if request.user.email else request.user.username
                software.updated_at = timezone.now()

            software.version_control += 1
            software.save()

            action_to = request.POST.getlist('action_to')
            finding_actions = FindingAction.objects.filter(finding=software)
            for finding_action in finding_actions:
                if finding_action.action_to not in action_to:
                    finding_action.delete()

            for email in action_to:

                finding_action, created = FindingAction.objects.get_or_create(
                    finding=software,
                    action_to=email
                )

                if finding_action.action_to and finding_action.status == EmailActionStatus.OPEN:
                    send_action_email(software, finding_action)

            prev_software = Finding.objects.create(
                **fields_and_values,
                user=request.user,
                version_control=software.version_control - 1
            )

            prev_software.parent = software
            prev_software.updated_by = software.updated_by
            prev_software.created_by = software.created_by
            prev_software.created_at = software.created_at
            prev_software.updated_at = timezone.now()
            prev_software.save()

            return redirect(reverse('finding'))

    db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]
    updated_db_field_names = [
        map_field_names[field] if field in map_field_names else field
        for field in db_field_names
    ]

    read_only_fields = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'selected_rows', 'action_status', 'deleted_at', 'deleted_by', 'created_at', 'updated_at', 'version_control', 'created_by', 'updated_by',)
    pre_column = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'selected_rows', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions',)
    compulsory_fields = ('status',)
    richtext_fields = ('description', 'recommendation', )
    finding_tabs = DependentDropdown.objects.filter(parent__title='Finding').values_list('title', flat=True)
    companies = DependentDropdown.objects.filter(parent__title='Finding')
    finding_categories = TableDropdownSubItem.objects.filter(item__item='Finding_Cat').values_list('subitem', flat=True)
    finding_questions = TableDropdownSubItem.objects.filter(item__item='Finding_Question').values_list('subitem', flat=True)
    action_type = TableDropdownSubItem.objects.filter(item__item='Action_Type').values_list('subitem', flat=True)
    
    COMBOS = {}
    COMBOS['action_to'] = User.objects.exclude(email="").values_list('email', flat=True)
    COMBOS['email_action_status'] = [choice[0] for choice in EmailActionStatus.choices]
    COMBOS['companies'] = finding_tabs
    COMBOS['itgc_categories'] = finding_categories
    COMBOS['itgc_questions'] = finding_questions
    COMBOS['action_type'] = action_type 

    form = FindingForm(initial={
        'description': software.description,
        'recommendation': software.recommendation
    })

    finding = Finding.objects.get(id=id)
    action_to_initial = finding.actions.values_list('action_to', flat=True)
    action_form = ActionForm(action_status=False, initial={'action_to': list(action_to_initial)})

    context = {
        'segment'  : 'finding',
        'parent'   : 'audit_report',
        'software': software,
        'db_field_names': updated_db_field_names,
        'read_only_fields': read_only_fields,
        'pre_column': pre_column,
        'date_picker_fields': Finding.date_fields_to_convert,
        'compulsory_fields': compulsory_fields,
        'richtext_fields': richtext_fields,
        'fav_name': 'AD Users',
        'int_fields': Finding.integer_fields,
        'float_fields': Finding.float_fields,
        'COMBOS': COMBOS,
        'form': form,
        'action_form': action_form,
        'companies': companies,
        'finding_url':  request.build_absolute_uri(f"{reverse('finding')}?finding_id={finding.id}")
    }
    
    return render(request, 'apps/edit/finding.html', context)



@login_required(login_url='/users/signin/')
def update_my_finding(request, id):
    software = Finding.objects.get(id=id)
    if request.method == 'POST':
        description = request.POST.get('description', '').strip()
        description_json = json.loads(description)
        description_html = description_json.get('html', '').strip()

        if not description_html or description_html == '<p><br></p>':
            messages.error(request, 'Description cannot be empty.')
            return redirect(request.META.get('HTTP_REFERER'))

        # fields = software._meta.get_fields()
        fields = [field.name for field in software._meta.get_fields() if not field.is_relation]

        fields_and_values = {}
        for field in fields:
            try:
                field_value = getattr(software, field)
                fields_and_values[field] = field_value
            except:
                continue

        fields_and_values.pop('id')
        # fields_and_values.pop('parent')
        fields_and_values.pop('created_at')
        fields_and_values.pop('created_by')
        fields_and_values.pop('updated_at')
        fields_and_values.pop('updated_by')
        fields_and_values.pop('version_control')

        if request.method == 'POST':
            for attribute, value in request.POST.items():
                if attribute == 'csrfmiddlewaretoken':
                    continue

                if value == '':
                    if attribute in Finding.integer_fields:
                        value = None
                    elif attribute in Finding.float_fields:
                        value = None
                    else:
                        continue

                if not attribute in Finding.date_fields_to_convert:
                    setattr(software, attribute, value)
                else:
                    unix_time = setattr(software, attribute, value)
                    possible_formats = ['%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M', '%Y-%m-%d %H:%M:%S', '%Y-%m-%d']
                    for format_str in possible_formats:
                        try:
                            if value:
                                unix_time = datetime.strptime(value, format_str).timestamp()
                                setattr(software, attribute, unix_time)
                                break
                        except ValueError:
                            pass 
            
            if software.action_status == 'DELETED':
                software.deleted_by = request.user.email if request.user.email else request.user.username
                software.deleted_at = timezone.now()
            else:
                software.updated_by = request.user.email if request.user.email else request.user.username
                software.updated_at = timezone.now()

            software.version_control += 1
            software.save()

            action_status = request.POST.get('user_action_status')
            finding_action = FindingAction.objects.filter(
                finding=software,
                action_to=request.user.email
            ).first()

            finding_action.status = action_status
            finding_action.save()

            if finding_action.action_to and finding_action.status == EmailActionStatus.CLOSE:
                send_action_email(software, finding_action)

            prev_software = Finding.objects.create(
                **fields_and_values,
                user=request.user,
                version_control=software.version_control - 1
            )

            prev_software.parent = software
            prev_software.updated_by = software.updated_by
            prev_software.created_by = software.created_by
            prev_software.created_at = software.created_at
            prev_software.updated_at = timezone.now()
            prev_software.save()

            return redirect(reverse('my_finding'))

    db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]
    updated_db_field_names = [
        map_field_names[field] if field in map_field_names else field
        for field in db_field_names
    ]
    read_only_fields = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by', 'created_at', 'updated_at', 'version_control', 'created_by', 'updated_by',)
    pre_column = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions','itgc_focus_area','risk','control',)
    compulsory_fields = ('status',)
    richtext_fields = ('description', 'recommendation', )
    finding_tabs = TableDropdownSubItem.objects.filter(item__item='Finding').values_list('subitem', flat=True)
    companies = DependentDropdown.objects.filter(parent__title='Finding')
    finding_categories = TableDropdownSubItem.objects.filter(item__item='Finding_Cat').values_list('subitem', flat=True)
    finding_questions = TableDropdownSubItem.objects.filter(item__item='Finding_Question').values_list('subitem', flat=True)
    action_type = TableDropdownSubItem.objects.filter(item__item='Action_Type').values_list('subitem', flat=True)
    
    COMBOS = {}
    COMBOS['action_to'] = User.objects.exclude(email="").values_list('email', flat=True)
    COMBOS['email_action_status'] = [choice[0] for choice in EmailActionStatus.choices]
    COMBOS['companies'] = finding_tabs
    COMBOS['itgc_categories'] = finding_categories
    COMBOS['itgc_questions'] = finding_questions
    COMBOS['action_type'] = action_type 

    form = FindingForm(initial={
        'description': software.description,
        'recommendation': software.recommendation
    })

    finding = Finding.objects.get(id=id)
    action_to_initial = finding.actions.values_list('action_to', flat=True)
    action_form = ActionForm(action_to=False, initial={'action_to': list(action_to_initial)})

    context = {
        'segment'  : 'finding',
        'parent'   : 'audit_report',
        'software': software,
        'db_field_names': updated_db_field_names,
        'read_only_fields': read_only_fields,
        'pre_column': pre_column,
        'date_picker_fields': Finding.date_fields_to_convert,
        'compulsory_fields': compulsory_fields,
        'richtext_fields': richtext_fields,
        'fav_name': 'AD Users',
        'int_fields': Finding.integer_fields,
        'float_fields': Finding.float_fields,
        'COMBOS': COMBOS,
        'form': form,
        'action_form': action_form,
        'companies': companies
    }
    
    return render(request, 'apps/edit/my-finding.html', context)


# Export as CSV
from django.utils.html import strip_tags
class ExportCSVView(View):
    def get(self, request, companies=None):
        db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]

        # Snapshot
        latest_snapshot, summary, snapshot_filter, snapshots = common_snapshot_filter(request, 'finding')

        fields = []
        pre_column = ('id', 'loader_instance', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions',)
        richtext_fields = ('description', 'recommendation', )
        dropdown_fields = ('companies', 'itgc_categories', 'itgc_questions', )

        # Show/Hide Filter Fields
        show_fields = HideShowFilter.objects.filter(value=False, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
        show_fields_keys = {field.key for field in show_fields}

        # Construct final fields list respecting original model order
        for field in db_field_names:
            # Ensure pre_column fields are not included
            if field not in pre_column and field in show_fields_keys:
                fields.append(field)

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="Actions.csv"'

        filter_string = {}
        filter_string['email_action_status__exact'] = 'OPEN'
        if companies:
            filter_string['companies__exact'] = companies

        filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.TEXT)
        combined_q_objects, user_unique_filter, user_query_conditions, user_count_filters = same_key_filter(filter_instance, return_count_filters=True)

        # for date range
        date_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.DATE)
        date_string, date_unique_filter, date_query_conditions, date_count_filters = common_date_filter(date_filter_instance, return_count_filters=True)
        filter_string.update(date_string)
        
        # for integer range
        int_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.INT)
        int_string, int_unique_filter, int_query_conditions, int_count_filters = common_integer_filter(int_filter_instance, return_count_filters=True)
        filter_string.update(int_string)

        # for float range
        float_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.FLOAT)
        float_string, float_unique_filter, float_query_conditions, float_count_filters = common_float_filter(float_filter_instance, return_count_filters=True)
        filter_string.update(float_string)

        order_by = request.GET.get('order_by', 'id')
        base_queryset = Finding.objects.filter(parent=None).filter(action_status=ActionStatus.IS_ACTIVE).filter(**filter_string).filter(combined_q_objects).filter(**snapshot_filter)
        queryset = base_queryset

        if user_query_conditions:
            queryset = queryset.filter(user_query_conditions)
        if date_query_conditions:
            queryset = queryset.filter(date_query_conditions)
        if int_query_conditions:
            queryset = queryset.filter(int_query_conditions)
        if float_query_conditions:
            queryset = queryset.filter(float_query_conditions)
        
        if user_count_filters:
            queryset = common_count_filter(user_count_filters, base_queryset, queryset, fields)
        elif date_count_filters:
            queryset = common_count_filter(date_count_filters, base_queryset, queryset, fields)
        elif int_count_filters:
            queryset = common_count_filter(int_count_filters, base_queryset, queryset, fields)
        elif float_count_filters:
            queryset = common_count_filter(float_count_filters, base_queryset, queryset, fields)
        else:
            if order_by == 'count' or order_by == '-count':
                order_by = 'ID'
            if 'count' in fields:
                fields.remove('count')

        queryset = queryset.order_by(order_by)

        if user_unique_filter:
            queryset = common_unique_filter(request, user_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if date_unique_filter:
            queryset = common_unique_filter(request, date_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if int_unique_filter:
            queryset = common_unique_filter(request, int_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if float_unique_filter:
            queryset = common_unique_filter(request, float_unique_filter, queryset, snapshot_filter, 'tables_finding')

        softwares = software_filter(request, queryset, db_field_names)

        writer = csv.writer(response)
        mapped_header_fields = [
            map_field_names.get(field, field)
            for field in fields
        ]
        mapped_header_fields = list(dict.fromkeys(mapped_header_fields))
        writer.writerow(mapped_header_fields) 

        for software in softwares:
            row_data = []
            for field in fields:
                mapped_field_name = reverse_map_field_names.get(field, field)
                if hasattr(software, mapped_field_name):
                    if mapped_field_name in Finding.date_fields_to_convert:
                        unix_timestamp = getattr(software, mapped_field_name)
                        if unix_timestamp:
                            date_time = datetime.utcfromtimestamp(unix_timestamp)
                            formatted_date = date_time.strftime('%b %d %Y %I:%M%p')
                            row_data.append(formatted_date)
                    elif mapped_field_name in richtext_fields:
                        attribute = getattr(software, mapped_field_name, None)
                        if attribute and hasattr(attribute, 'html'):
                            row_data.append(attribute.html)
                        else:
                            row_data.append("")
                    elif mapped_field_name in dropdown_fields:
                        attribute = getattr(software, mapped_field_name, None)
                        try:
                            value = int(attribute)
                            dropdown = DependentDropdown.objects.filter(pk=value).first()
                            if dropdown:
                                row_data.append(dropdown.title)
                            else:
                                row_data.append("")
                        except:
                            row_data.append("")
                    else:
                        row_data.append(getattr(software, mapped_field_name))
                
                else:
                    row_data.append("")

            writer.writerow(row_data)

        return response


class ExportMyFindingCSVView(View):
    def get(self, request, companies=None):
        action_status = request.GET.get('action_status')
        db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]
        action_filter = {}
        action_filter['status__exact'] = 'OPEN'
        if action_status:
            if not action_status == 'all':
                action_filter['status__exact'] = action_status
            else:
                action_filter = {}

        finding_ids = FindingAction.objects.filter(**action_filter).filter(action_to=request.user.email).values_list('finding__id', flat=True)

        # Snapshot
        latest_snapshot, summary, snapshot_filter, snapshots = common_snapshot_filter(request, 'finding')

        fields = []
        pre_column = ('id', 'loader_instance', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions',)
        richtext_fields = ('description', 'recommendation', )
        dropdown_fields = ('companies', 'itgc_categories', 'itgc_questions', )

        # Show/Hide Filter Fields
        show_fields = HideShowFilter.objects.filter(value=False, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
        show_fields_keys = {field.key for field in show_fields}

        # Construct final fields list respecting original model order
        for field in db_field_names:
            # Ensure pre_column fields are not included
            if field not in pre_column and field in show_fields_keys:
                fields.append(field)

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="Actions.csv"'

        filter_string = {}
        filter_string['email_action_status__exact'] = 'OPEN'

        filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.TEXT)
        combined_q_objects, user_unique_filter, user_query_conditions, user_count_filters = same_key_filter(filter_instance, return_count_filters=True)

        # for date range
        date_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.DATE)
        date_string, date_unique_filter, date_query_conditions, date_count_filters = common_date_filter(date_filter_instance, return_count_filters=True)
        filter_string.update(date_string)

        # for integer range
        int_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.INT)
        int_string, int_unique_filter, int_query_conditions, int_count_filters = common_integer_filter(int_filter_instance, return_count_filters=True)
        filter_string.update(int_string)

        # for float range
        float_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.FLOAT)
        float_string, float_unique_filter, float_query_conditions, float_count_filters = common_float_filter(float_filter_instance, return_count_filters=True)
        filter_string.update(float_string)

        order_by = request.GET.get('order_by', 'id')
        base_queryset = Finding.objects.filter(id__in=finding_ids).filter(parent=None).filter(action_status=ActionStatus.IS_ACTIVE).filter(**filter_string).filter(combined_q_objects).filter(**snapshot_filter)
        queryset = base_queryset

        if user_query_conditions:
            queryset = queryset.filter(user_query_conditions)
        if date_query_conditions:
            queryset = queryset.filter(date_query_conditions)
        if int_query_conditions:
            queryset = queryset.filter(int_query_conditions)
        if float_query_conditions:
            queryset = queryset.filter(float_query_conditions)

        if user_count_filters:
            queryset = common_count_filter(user_count_filters, base_queryset, queryset, fields)
        elif date_count_filters:
            queryset = common_count_filter(date_count_filters, base_queryset, queryset, fields)
        elif int_count_filters:
            queryset = common_count_filter(int_count_filters, base_queryset, queryset, fields)
        elif float_count_filters:
            queryset = common_count_filter(float_count_filters, base_queryset, queryset, fields)
        else:
            if order_by == 'count' or order_by == '-count':
                order_by = 'ID'
            if 'count' in fields:
                fields.remove('count')

        queryset = queryset.order_by(order_by)

        if user_unique_filter:
            queryset = common_unique_filter(request, user_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if date_unique_filter:
            queryset = common_unique_filter(request, date_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if int_unique_filter:
            queryset = common_unique_filter(request, int_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if float_unique_filter:
            queryset = common_unique_filter(request, float_unique_filter, queryset, snapshot_filter, 'tables_finding')

        softwares = software_filter(request, queryset, db_field_names)

        writer = csv.writer(response)
        mapped_header_fields = [
            map_field_names.get(field, field)
            for field in fields
        ]
        mapped_header_fields = list(dict.fromkeys(mapped_header_fields))
        writer.writerow(mapped_header_fields) 

        for software in softwares:
            row_data = []
            for field in fields:
                mapped_field_name = reverse_map_field_names.get(field, field)
                if hasattr(software, mapped_field_name):
                    if mapped_field_name in Finding.date_fields_to_convert:
                        unix_timestamp = getattr(software, mapped_field_name)
                        if unix_timestamp:
                            date_time = datetime.utcfromtimestamp(unix_timestamp)
                            formatted_date = date_time.strftime('%b %d %Y %I:%M%p')
                            row_data.append(formatted_date)
                    elif mapped_field_name in richtext_fields:
                        attribute = getattr(software, mapped_field_name, None)
                        if attribute and hasattr(attribute, 'html'):
                            row_data.append(attribute.html)
                        else:
                            row_data.append("")
                    elif mapped_field_name in dropdown_fields:
                        attribute = getattr(software, mapped_field_name, None)
                        try:
                            value = int(attribute)
                            dropdown = DependentDropdown.objects.filter(pk=value).first()
                            if dropdown:
                                row_data.append(dropdown.title)
                            else:
                                row_data.append("")
                        except:
                            row_data.append("")
                    else:
                        row_data.append(getattr(software, mapped_field_name))

                else:
                    row_data.append("")

            writer.writerow(row_data)

        return response

from reportlab.pdfgen import canvas
from io import BytesIO
# Export as PDF
class ExportPDFView(View):
    def get(self, request):
        buffer = BytesIO()
        response_pdf = HttpResponse(content_type='application/pdf')
        response_pdf['Content-Disposition'] = 'attachment; filename="softwares.pdf"'

        pdf = canvas.Canvas(buffer)

        page_number = request.GET.get('page', 1)
        items_per_page = 25

        softwares = Finding.objects.all()
        paginator = Paginator(softwares, items_per_page)
        current_page_devices = paginator.get_page(page_number)

        y_position = 800

        for device in current_page_devices:
            pdf.drawString(100, y_position, f"id: {device.id}, DeviceName: {device.DeviceName}, OperatingSystem: {device.OperatingSystem}")
            y_position -= 20

        pdf.save()

        buffer.seek(0)
        response_pdf.write(buffer.getvalue())

        return response_pdf



def finding_version_control_dt(request, id):
    parent = Finding.objects.get(id=id)
    db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]
    updated_db_field_names = [
        map_field_names[field] if field in map_field_names else field
        for field in db_field_names
    ]
    pre_column = ('id', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions',)
    richtext_fields = ('description', 'recommendation', )
    dropdown_fields = ('companies', 'itgc_categories', 'itgc_questions', )
    dropdown_fields = [
        map_field_names[field] if field in map_field_names else field
        for field in dropdown_fields
    ]


    order_by = request.GET.get('order_by', '-version_control')
    child_qs = Finding.objects.filter(parent=parent).order_by(order_by)

    context = {
        'segment'  : 'finding',
        'parent'   : 'audit_report',
        # 'form'     : form,
        'softwares' : child_qs,
        'db_field_names': updated_db_field_names,
        'pre_column': pre_column,
        'richtext_fields': richtext_fields,
        'dropdown_fields': dropdown_fields,
    }

    return render(request, 'apps/version_control/finding.html', context)




# Inline export
import re
import base64
from apps.tables.models import Finding
from docx import Document
from htmldocx import HtmlToDocx
from docx.shared import Inches

class ExportDocxView(View):
    def get(self, request, companies=None):
        document = Document()
        new_parser = HtmlToDocx()

        db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]

        # Snapshot
        latest_snapshot, summary, snapshot_filter, snapshots = common_snapshot_filter(request, 'finding')

        fields = []
        pre_column = ('id', 'loader_instance', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions',)
        richtext_fields = ('description', 'recommendation', )
        dropdown_fields = ('companies', 'itgc_categories', 'itgc_questions', )

        show_fields = HideShowFilter.objects.filter(value=False, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
        for field in show_fields:
            if field.key not in pre_column:
                if field.key not in fields:
                    fields.append(field.key)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Actions.docx"'

        filter_string = {}
        filter_string['email_action_status__exact'] = 'OPEN'
        if companies:
            filter_string['companies__exact'] = companies
        filter_instance = UserFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
        combined_q_objects, user_unique_filter, user_query_conditions, user_count_filters = same_key_filter(filter_instance, return_count_filters=True)

        # for date range
        date_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.DATE)
        date_string, date_unique_filter, date_query_conditions, date_count_filters = common_date_filter(date_filter_instance, return_count_filters=True)
        filter_string.update(date_string)
        
        # for integer range
        int_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.INT)
        int_string, int_unique_filter, int_query_conditions, int_count_filters = common_integer_filter(int_filter_instance, return_count_filters=True)
        filter_string.update(int_string)

        # for float range
        float_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.FLOAT)
        float_string, float_unique_filter, float_query_conditions, float_count_filters = common_float_filter(float_filter_instance, return_count_filters=True)
        filter_string.update(float_string)

        order_by = request.GET.get('order_by', 'id')
        base_queryset = Finding.objects.filter(parent=None).filter(action_status=ActionStatus.IS_ACTIVE).filter(**filter_string).filter(combined_q_objects).filter(**snapshot_filter)
        queryset = base_queryset

        if user_query_conditions:
            queryset = queryset.filter(user_query_conditions)
        if date_query_conditions:
            queryset = queryset.filter(date_query_conditions)
        if int_query_conditions:
            queryset = queryset.filter(int_query_conditions)
        if float_query_conditions:
            queryset = queryset.filter(float_query_conditions)
        
        if user_count_filters:
            queryset = common_count_filter(user_count_filters, base_queryset, queryset, fields)
        elif date_count_filters:
            queryset = common_count_filter(date_count_filters, base_queryset, queryset, fields)
        elif int_count_filters:
            queryset = common_count_filter(int_count_filters, base_queryset, queryset, fields)
        elif float_count_filters:
            queryset = common_count_filter(float_count_filters, base_queryset, queryset, fields)
        else:
            if order_by == 'count' or order_by == '-count':
                order_by = 'ID'
            if 'count' in fields:
                fields.remove('count')

        queryset = queryset.order_by(order_by)

        if user_unique_filter:
            queryset = common_unique_filter(request, user_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if date_unique_filter:
            queryset = common_unique_filter(request, date_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if int_unique_filter:
            queryset = common_unique_filter(request, int_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if float_unique_filter:
            queryset = common_unique_filter(request, float_unique_filter, queryset, snapshot_filter, 'tables_finding')

        softwares = software_filter(request, queryset, db_field_names)

        def extract_base64_images(html_content):
            base64_images = re.findall(r'data:image/(png|jpeg);base64,([a-zA-Z0-9+/=]+)', html_content)
            return base64_images, re.sub(r'data:image/(png|jpeg);base64,([a-zA-Z0-9+/=]+)', '', html_content)

        for idx, software in enumerate(softwares, start=1):
            document.add_heading(f'Finding {idx}', level=1)

            for field in fields:
                mapped_field_name = reverse_map_field_names.get(field, field)
                if hasattr(software, mapped_field_name):
                    value = getattr(software, mapped_field_name, "")
                    if mapped_field_name in Finding.date_fields_to_convert:
                        unix_timestamp = getattr(software, mapped_field_name)
                        if unix_timestamp:
                            date_time = datetime.utcfromtimestamp(unix_timestamp)
                            formatted_date = date_time.strftime('%b %d %Y %I:%M%p')
                            document.add_paragraph(f'{mapped_field_name}: {formatted_date}')
                    elif mapped_field_name in richtext_fields:
                        attribute = getattr(software, mapped_field_name, None)
                        document.add_paragraph(f'{mapped_field_name.capitalize()}:')
                        if attribute and hasattr(attribute, 'html'):
                            html_content = attribute.html
                            base64_images, cleaned_html = extract_base64_images(html_content)

                            new_parser.add_html_to_document(cleaned_html, document)

                            for img_format, img_data in base64_images:
                                img_binary = base64.b64decode(img_data)
                                image_stream = BytesIO(img_binary)
                                document.add_picture(image_stream, width=Inches(4))
                        else:
                            document.add_paragraph(f'{mapped_field_name}: [No Content]')
                    elif mapped_field_name in dropdown_fields:
                        attribute = getattr(software, mapped_field_name, None)
                        mapped_field = map_field_names.get(mapped_field_name, mapped_field_name)
                        try:
                            value = int(attribute)
                            dropdown = DependentDropdown.objects.filter(pk=value).first()
                            if dropdown:
                                document.add_paragraph(f'{mapped_field.capitalize()}: {dropdown.title}')
                            else:
                                document.add_paragraph(f'{mapped_field.capitalize()}: ""')
                        except:
                            document.add_paragraph(f'{mapped_field.capitalize()}: ""')
                    else:
                        document.add_paragraph(f'{mapped_field_name.capitalize()}: {value}')


        # Specify the string to remove (e.g., "<image: >")
        target_string = "<image: >"
        # Call the function to remove the string
        document = remove_string_from_document(document, target_string)
        # Save the modified document
        document.save(response)
        return response
    
class ExportMyFindingDocxView(View):
    def get(self, request, companies=None):
        document = Document()
        new_parser = HtmlToDocx()

        action_status = request.GET.get('action_status')
        db_field_names = [field.name for field in Finding._meta.get_fields() if not field.is_relation]
        action_filter = {}
        action_filter['status__exact'] = 'OPEN'
        if action_status:
            if not action_status == 'all':
                action_filter['status__exact'] = action_status
            else:
                action_filter = {}

        finding_ids = FindingAction.objects.filter(**action_filter).filter(action_to=request.user.email).values_list('finding__id', flat=True)

        # Snapshot
        latest_snapshot, summary, snapshot_filter, snapshots = common_snapshot_filter(request, 'finding')

        fields = []
        pre_column = ('id', 'loader_instance', 'user', 'name', 'content_type', 'model_choices', 'pre_filters', 'pre_columns', 'richtext_fields', 'page_items', 'hide_show_filters', 'user_filters', 'server_filters', 'date_range_filters', 'int_range_filters', 'float_range_filters', 'search', 'order_by', 'snapshot', 'query_snapshot', 'is_dynamic_query', 'has_documents', 'is_split_dt', 'parent_dt', 'match_field', 'child_dt', 'action_status', 'deleted_at', 'deleted_by','companies', 'itgc_categories', 'itgc_questions',)
        richtext_fields = ('description', 'recommendation', )
        dropdown_fields = ('companies', 'itgc_categories', 'itgc_questions', )

        show_fields = HideShowFilter.objects.filter(value=False, userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
        for field in show_fields:
            if field.key not in pre_column:
                if field.key not in fields:
                    fields.append(field.key)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Actions.docx"'

        filter_string = {}
        filter_string['email_action_status__exact'] = 'OPEN'
        filter_instance = UserFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW)
        combined_q_objects, user_unique_filter, user_query_conditions, user_count_filters = same_key_filter(filter_instance, return_count_filters=True)

        # for date range
        date_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.DATE)
        date_string, date_unique_filter, date_query_conditions, date_count_filters = common_date_filter(date_filter_instance, return_count_filters=True)
        filter_string.update(date_string)

        # for integer range
        int_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.INT)
        int_string, int_unique_filter, int_query_conditions, int_count_filters = common_integer_filter(int_filter_instance, return_count_filters=True)
        filter_string.update(int_string)

        # for float range
        float_filter_instance = SavedFilter.objects.filter(userID=get_user_id(request), parent=ModelChoices.FINDING_VIEW, field_type=FieldType.FLOAT)
        float_string, float_unique_filter, float_query_conditions, float_count_filters = common_float_filter(float_filter_instance, return_count_filters=True)
        filter_string.update(float_string)

        order_by = request.GET.get('order_by', 'id')
        base_queryset = Finding.objects.filter(id__in=finding_ids).filter(parent=None).filter(action_status=ActionStatus.IS_ACTIVE).filter(**filter_string).filter(combined_q_objects).filter(**snapshot_filter)
        queryset = base_queryset

        if user_query_conditions:
            queryset = queryset.filter(user_query_conditions)
        if date_query_conditions:
            queryset = queryset.filter(date_query_conditions)
        if int_query_conditions:
            queryset = queryset.filter(int_query_conditions)
        if float_query_conditions:
            queryset = queryset.filter(float_query_conditions)

        if user_count_filters:
            queryset = common_count_filter(user_count_filters, base_queryset, queryset, fields)
        elif date_count_filters:
            queryset = common_count_filter(date_count_filters, base_queryset, queryset, fields)
        elif int_count_filters:
            queryset = common_count_filter(int_count_filters, base_queryset, queryset, fields)
        elif float_count_filters:
            queryset = common_count_filter(float_count_filters, base_queryset, queryset, fields)
        else:
            if order_by == 'count' or order_by == '-count':
                order_by = 'ID'
            if 'count' in fields:
                fields.remove('count')

        queryset = queryset.order_by(order_by)

        if user_unique_filter:
            queryset = common_unique_filter(request, user_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if date_unique_filter:
            queryset = common_unique_filter(request, date_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if int_unique_filter:
            queryset = common_unique_filter(request, int_unique_filter, queryset, snapshot_filter, 'tables_finding')
        if float_unique_filter:
            queryset = common_unique_filter(request, float_unique_filter, queryset, snapshot_filter, 'tables_finding')

        softwares = software_filter(request, queryset, db_field_names)

        def extract_base64_images(html_content):
            base64_images = re.findall(r'data:image/(png|jpeg);base64,([a-zA-Z0-9+/=]+)', html_content)
            return base64_images, re.sub(r'data:image/(png|jpeg);base64,([a-zA-Z0-9+/=]+)', '', html_content)

        for idx, software in enumerate(softwares, start=1):
            document.add_heading(f'Finding {idx}', level=1)

            for field in fields:
                mapped_field_name = reverse_map_field_names.get(field, field)
                if hasattr(software, mapped_field_name):
                    value = getattr(software, mapped_field_name, "")
                    if mapped_field_name in Finding.date_fields_to_convert:
                        unix_timestamp = getattr(software, mapped_field_name)
                        if unix_timestamp:
                            date_time = datetime.utcfromtimestamp(unix_timestamp)
                            formatted_date = date_time.strftime('%b %d %Y %I:%M%p')
                            document.add_paragraph(f'{mapped_field_name}: {formatted_date}')
                    elif mapped_field_name in richtext_fields:
                        attribute = getattr(software, mapped_field_name, None)
                        document.add_paragraph(f'{mapped_field_name.capitalize()}:')
                        if attribute and hasattr(attribute, 'html'):
                            html_content = attribute.html
                            base64_images, cleaned_html = extract_base64_images(html_content)

                            new_parser.add_html_to_document(cleaned_html, document)

                            for img_format, img_data in base64_images:
                                img_binary = base64.b64decode(img_data)
                                image_stream = BytesIO(img_binary)
                                document.add_picture(image_stream, width=Inches(4))
                        else:
                            document.add_paragraph(f'{mapped_field_name}: [No Content]')
                    elif mapped_field_name in dropdown_fields:
                        attribute = getattr(software, mapped_field_name, None)
                        mapped_field = map_field_names.get(mapped_field_name, mapped_field_name)
                        try:
                            value = int(attribute)
                            dropdown = DependentDropdown.objects.filter(pk=value).first()
                            if dropdown:
                                document.add_paragraph(f'{mapped_field.capitalize()}: {dropdown.title}')
                            else:
                                document.add_paragraph(f'{mapped_field.capitalize()}: ""')
                        except:
                            document.add_paragraph(f'{mapped_field.capitalize()}: ""')
                    else:
                        document.add_paragraph(f'{mapped_field_name.capitalize()}: {value}')


        # Specify the string to remove (e.g., "<image: >")
        target_string = "<image: >"
        # Call the function to remove the string
        document = remove_string_from_document(document, target_string)
        # Save the modified document
        document.save(response)
        return response

def remove_string_from_document(document, target_string):
    # Iterate through all paragraphs in the document
    for paragraph in document.paragraphs:
        # Access the underlying XML of the paragraph
        for run in paragraph.runs:
            # Modify the text in each run if the target string is found
            if target_string in run.text:
                run.text = run.text.replace(target_string, "")

    return document

